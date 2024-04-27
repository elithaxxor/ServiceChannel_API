import base64, os, json, re, datetime
import requests
from requests_oauthlib import OAuth2Session
from datetime import time
import pprint
from docx import Document
from docx.shared import Inches

'''
TODO LIST: 
    def json2dict(toParse) --> Finish the function where the document is built out
'''

''' 
THE PROGRAM RUNS A POST REQUEST TO THE URL, 'https://sb2login.servicechannel.com/oauth/token'
THE PROGRAM SAVES THE TOKENS (ACCESS AND REFRESH) IN A GLOBAL DICTIONARY, 
CALLED--  TOKEN_TABLE, ACCESS_TOKEN, REFRESH_TOKEN 
THE PROGRAM THEN USES THE TOKENS TO ACCESS THE API, AND FETCHES THE WORKORDER JSON FILE
IT THEN WRITES THE JSON FILE TO A LOCAL FILE, 'workorder{workOrderId}.json'

'''
# api_key = "8CA8C50D-2A0C-46DD-8653-23F7D32F6D2A"
url = "https://sb2login.servicechannel.com/oauth/token"
url2 = "https://login.servicechannel.com/oauth/token"


authorization_endpoint = "https://login.servicechannel.com/oauth/authorize" # FOR SAND BOX ENVIROMENT
token_endpoint = "https://login.servicechannel.com/oauth/token" # FOR PRODUCTION ENVIROMEN T
token_endpoint2 = "https://api.servicechannel.com/" # API GATEWAY
client_id = "PR.2098457598.54320AAF-C2F9-447E-BE61-8C22D96CCE7F" # SPECIFIC TO
client_secret = '8CA8C50D-2A0C-46DD-8653-23F7D32F6D2A'

''' I encoded 2 application keys, try both  '''
#auth_header = "Basic UFIuMjA5ODQ1NzU5OC41NDMyMEFBRi1DMkY5LTQ0N0UtQkU2MS04QzIyRDk2Q0NFN0Y6OENBOEM1MEQtMkEwQy00NkRELTg2NTMtMjNGN0QzMkY2RDJB"
#auth_header = "Basic UFIuMjA5ODQ1NzU5OC5GNjUxMEYxOS03MTExLTQ4NjMtQTQwMi0yQ0Y4OTlGQTNGOUI6ODlDODg3RTctMEY1My00MUE5LUI4OTUtMjg2MzUxNzdFQTY4"
auth_header = "Basic UFIuMjA5ODQ1NzU5OC4yNkFBMkVFOS0xMTUwLTRDQzctOEU0Qi1ENzk2MUM3NDFBRDY6RjQ4RDcxRDItNEFEOS00RjQ0LUE0MTktRjMxMUFBMUNEQTk3"


''' TAKES THE TOKEN JSON RESPONSE AND PARSES IT INTO A GLOBAL PYTHON DICTIONARY '''
def MAKE_TOKEN_GLOBAL(toParse):

    print('[!] Running JSON Parser')
    global TOKEN_TABLE, ACCESS_TOKEN, REFRESH_TOKEN # GLOBAL TABLE TO ACCESS DATA AT A LATER TIME
    TOKEN_TABLE = json.loads(toParse)
    ACCESS_TOKEN = TOKEN_TABLE['access_token']
    REFRESH_TOKEN = TOKEN_TABLE['refresh_token']

    print('[!] Token Tabel', TOKEN_TABLE)

import json

class TokenParser:
    def __init__(self, toParse):
        self.token_table = json.loads(toParse)
        self.access_token = self.token_table['access_token']
        self.refresh_token = self.token_table['refresh_token']

    def get_token_table(self):
        return self.token_table

    def get_access_token(self):
        return self.access_token

    def get_refresh_token(self):
        return self.refresh_token



def get_resp():
    # HEADER FOR : /oauth/token
    ''' 'username': 'fuad@dedicatedglass.com',
        'password': 'Dedicated!234',
        'grant_type': 'password',
    '''

    headers = {
        'Authorization': auth_header,
        'Content-Type': "application/x-www-form-urlencoded"
    }

    payload = {
        "username": "fuad@dedicatedglass.com",
        "password": "Dedicated!234",
        "grant_type": "password"
    }

    # Send the POST request
    response = requests.post(url, headers=headers, data=payload)
    print("[+]\n" ,response.text)
    print("[+]\n" , response.json)
    print("[+]\n" , response.status_code)

    if response.status_code == 200:
        # data = response.json()
        new_tokens = response.json()
        global access_token
        global refresh_token
        access_token = new_tokens['access_token']
        refresh_token = new_tokens['refresh_token']
        respone_header = response.headers

        print("[+]\n", response.text)
        print('[+] TOKEN: \n', access_token)
        print('[+] REFRESH TOKEN: \n', refresh_token)
        print('[+] EXPIRES IN \n', new_tokens['expires_in'])

       # print("[!] RESPONSE HEADERS FROM GET RESP")
        # for (key, value) in respone_header.items():
        #     print(f"{key}: {value}")

        MAKE_TOKEN_GLOBAL(response.text)
        ## TOKEN RECEIVED, ACCES API

        accessAPI(access_token, refresh_token)

    else:
        print(f"Failed to authenticate: {response.status_code} - {response.content}")
        print(response.reason)

def accessAPI(access_token, refresh_token):
    print("[!] Running 'ACCESS API' \n [ACCESS TOKEN]: ", access_token)
    print("[REFRESH TOKEN]: ", refresh_token)

    ## ACCESSES GLOBAL DICTIONARY, WHERE TOKENS ARE STORED FROM JSON RESPONSE

    workOrderId = 273890369
    url = f"https://sb2api.servicechannel.com/v3/workorders/{workOrderId}"
    base_url = "api.servicechannel.com"

    # Define the URL
    #url = "https://api.servicechannel.com/v3/workorders/1000000"
    auth_header = f"Bearer {access_token}"

    # Define the headers
    headers = {
        "Authorization": auth_header,
        "Content-Type": "application/x-www-form-urlencoded"
    }


    # Send the POST request
    response = requests.get(url, headers=headers)
    respone_header = response.headers

    # print("[!] RESPONSE HEADERS FROM APIACCESS")
    # for (key, value) in respone_header.items():
    #     print(f"{key}: {value}")


    ## JSON FORMATTING
    resp_json = response.json()
    resp_str = response.content
    #json2dict(resp_str)

    # print json to screen with human-friendly formatting
    pprint.pprint(resp_json, compact=True)

    # write json to file with human-friendly formatting
    pretty_json_str = pprint.pformat(resp_json, compact=True).replace("'", '"')
    with open(f'workorder{workOrderId}.json', 'w') as f:
        f.write(pretty_json_str)

    # Print the response
    print("[!] RESPONSE TEXT \n", response.text)
    print("[!] RESPONSE JSON \n", response.json)
    print("[!] RESPONSE STATUS \n", response.status_code)

def refreshToken(access_token, refresh_token):
    print("[!] Running 'REFRESH TOKEN' \n [ACCESS TOKEN]: ", access_token)
    print("[REFRESH TOKEN]: ", refresh_token)

    # Define the headers
    auth_header = f"Basic {access_token}"
    headers = {
        "Authorization": auth_header,
        'Content-Type': "application/x-www-form-urlencoded"
    }

    # Define the parameters
    params = {
        "refresh_token": refresh_token,
        "grant_type": "refresh_token"
    }

    # Send the POST request
    response = requests.post(url, headers=headers, data=params)

    # Print the response
    print("[!] RESPONSE TEXT \n", response.text)
    print("[!] RESPONSE JSON \n", response.json)
    print("[!] RESPONSE STATUS \n", response.status_code)

''' TAKES THE JSON WORKORDER RESPONSE AND PARSES IT INTO A PYTHON DICTIONARY '''
def json2dict(toParse):
    # Parse JSON string into Python dictionary
    workorder = json.loads(toParse)

    # Assign values to variables with the keys
    ApprovalCode = workorder['ApprovalCode']
    AssetCount = workorder['AssetCount']
    Attachments = workorder['Attachments']
    AutoComplete = workorder['AutoComplete']
    AutoInvoice = workorder['AutoInvoice']
    CallDate = workorder['CallDate']
    CallDate_DTO = workorder['CallDate_DTO']
    Caller = workorder['Caller']
    CallerId = workorder['CallerId']
    Category = workorder['Category']
    CategoryId = workorder['CategoryId']
    CheckInDeniedReason = workorder['CheckInDeniedReason']
    CheckInRange = workorder['CheckInRange']
    CreatedBy = workorder['CreatedBy']
    Currency = workorder['Currency']
    Description = workorder['Description']
    ExpirationDate = workorder['ExpirationDate']
    ExpirationDate_DTO = workorder['ExpirationDate_DTO']
    HasWorkActivity = workorder['HasWorkActivity']
    Id = workorder['Id']
    IsCheckInDenied = workorder['IsCheckInDenied']
    IsEnabledForMobile = workorder['IsEnabledForMobile']
    IsExpired = workorder['IsExpired']
    IsInvoiced = workorder['IsInvoiced']
    IssueTicketInfo = workorder['IssueTicketInfo']
    Labels = workorder['Labels']
    Location = workorder['Location']
    LocationId = workorder['LocationId']
    Notes = workorder['Notes']
    Nte = workorder['Nte']
    Number = workorder['Number']
    PostedId = workorder['PostedId']
    Priority = workorder['Priority']
    ProblemCode = workorder['ProblemCode']
    Provider = workorder['Provider']
    PurchaseNumber = workorder['PurchaseNumber']
    Resolution = workorder['Resolution']
    ScheduledDate = workorder['ScheduledDate']
    ScheduledDate_DTO = workorder['ScheduledDate_DTO']
    Source = workorder['Source']
    Status = workorder['Status']
    Subscriber = workorder['Subscriber']
    Trade = workorder['Trade']
    UpdatedDate = workorder['UpdatedDate']
    UpdatedDate_DTO = workorder['UpdatedDate_DTO']

    def makeDoc():
        document = Document()
        document.add_heading('workOrder : ', 0)
        p = document.add_paragraph('A plain paragraph having some ')

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        document.save('demo.docx')

    # Print the variables
    print(f'ApprovalCode: {ApprovalCode}')
    print(f'AssetCount: {AssetCount}')
    print(f'Attachments: {Attachments}')
    print(f'AutoComplete: {AutoComplete}')
    print(f'AutoInvoice: {AutoInvoice}')
    print(f'CallDate: {CallDate}')
    print(f'CallDate_DTO: {CallDate_DTO}')
    print(f'Caller: {Caller}')
    print(f'CallerId: {CallerId}')
    print(f'Category: {Category}')
    print(f'CategoryId: {CategoryId}')
    print(f'CheckInDeniedReason: {CheckInDeniedReason}')
    print(f'CheckInRange: {CheckInRange}')
    print(f'CreatedBy: {CreatedBy}')
    print(f'Currency: {Currency}')
    print(f'Description: {Description}')
    print(f'ExpirationDate: {ExpirationDate}')
    print(f'ExpirationDate_DTO: {ExpirationDate_DTO}')
    print(f'HasWorkActivity: {HasWorkActivity}')
    print(f'Id: {Id}')
    print(f'IsCheckInDenied: {IsCheckInDenied}')
    print(f'IsEnabledForMobile: {IsEnabledForMobile}')
    print(f'IsExpired: {IsExpired}')
    print(f'IsInvoiced: {IsInvoiced}')
    print(f'IssueTicketInfo: {IssueTicketInfo}')
    print(f'Labels: {Labels}')
    print(f'Location: {Location}')
    print(f'LocationId: {LocationId}')
    print(f'Notes: {Notes}')
    print(f'Nte: {Nte}')
    print(f'Number: {Number}')
    print(f'PostedId: {PostedId}')
    print(f'Priority: {Priority}')
    print(f'ProblemCode: {ProblemCode}')
    print(f'Provider: {Provider}')
    print(f'PurchaseNumber: {PurchaseNumber}')
    print(f'Resolution: {Resolution}')
    print(f'ScheduledDate: {ScheduledDate}')
    print(f'ScheduledDate_DTO: {ScheduledDate_DTO}')
    print(f'Source: {Source}')
    print(f'Status: {Status}')
    print(f'Subscriber: {Subscriber}')
    print(f'Trade: {Trade}')
    print(f'UpdatedDate: {UpdatedDate}')
    print(f'UpdatedDate_DTO: {UpdatedDate_DTO}')




if __name__ == '__main__':

    get_resp()
   # runOauth()
   #runSandBox()






def runSandBox():
    # Define the headers
    headers = {
        "Content-Type": "application/x-www-form-urlencoded",
        "Authorization": "Basic U0IuMjAxNDkxNzI0My43MDk0MzU1MC05OTc0LTQxRTctOEIzOC1DQThDOEMwMjQyNjI6MUM2MDBDQjctNEQyNC00Nzg3LUJCM0UtQUNEMkIyRkJGNzQ1"
    }

    # Define the body parameters
    body = {
        "username": "SC-Dev1",
        "password": "servicechannel1",
        "grant_type": "password"
    }

    # Send the POST request
    response = requests.post(url, headers=headers, data=body)

    # Print the response
    print(response.json())


def runOauth():
    # Create an OAuth2 session
    oauth2_session = OAuth2Session(client_id, redirect_uri="")

    # Generate the authorization URL
    authorization_url, state = oauth2_session.authorization_url(token_endpoint)

    # Print the authorization URL and ask the user to visit it
    print(f"Please visit the following URL to authorize the application:\n{authorization_url}")

    # Get the authorization code from the user (after successful authorization)
    authorization_code = input("Enter the authorization code: ")

    # Fetch the access token using the authorization code
    token = oauth2_session.fetch_token(
        token_endpoint,
        authorization_response=authorization_code,
        client_secret=client_secret
    )

    # Print the access token
    print(f"Access token: {token['access_token']}")

    # print(response.text)


