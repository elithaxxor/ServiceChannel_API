import logging, requests, json, html, time, os, sys, urllib.request
from flask import Flask, request, jsonify, json, Blueprint, render_template, request, jsonify, redirect, url_for,render_template_string, make_response, send_file, flash
import pprint as pp
import main as main
from flask_socketio import SocketIO, emit
from wtforms import FileField, SubmitField
from flask_wtf import FlaskForm
from docx import Document
from werkzeug.utils import secure_filename
from wtforms.validators import InputRequired


'''
    [USAGE]: 
    * Run Flask: python app.py
        https://33b6-24-44-49-28.ngrok-free.app/ (root) 
        https://33b6-24-44-49-28.ngrok-free.app/web/jsonobj  (auto update)
        
    [NOTES] 
        * EVENT DRIVEN METHOLOGY.
            + WebHooks - A way for an app to provide other applications with real-time information.
                - When an event occurs that the user is subscribed to (e.g., a new message), the app sends an HTTP request to the URL configured by the user.
            + WebSockets - A communication protocol that provides full-duplex communication channels over a single TCP connection.
                - The server can push data to the client without the client needing to request it.
            + Long Polling - A web application sends a request to the server and the server holds the request open until new data is available.
                - The server sends a response to the client, which immediately sends a new request.
        
        [UPLOADING FILES]
        * THE USER UPLOADS FILES USING FLASKS WTFORMS FLASK_WTF AND WEKZEUG UTILS
            -> WTForms is a flexible forms validation and rendering library for Python web development.
            -> Flask-WTF is a Flask extension that makes working with forms easier.
            -> Werkzeug is a WSGI utility library for Python.
            -> Secure_filename: Returns a secure version of the filename.
            -> FileField: Creates a file upload field.
            -> SubmitField: Creates a submit button.
            -> FlaskForm: A class that represents a form.
            -> send_file: Sends a file to the client.
            
        [DOWNLOADING FILES]
        * THE USER DOWNLOADS FILES USING A JAVASCRIPT FUNCTION AND BUTTON THAT FETCHES THE FILE FROM THE SERVER
        
    [STORAGE] -- 
        * [FUNCTION] StoreData: Stores / retrieves data from the webhook
        * [FUNCTION] storeToken: Stores / retrieves token from the API 
        * [FUNCTION] makeCustomHeaders: Creates custom headers for the API
        * [FUNCTION] TokenParser: Parses the token from the API and stores it in a dictionary
        * [FUNCTION] JSONDocumentCreator: Creates a document from JSON data
        * [FUNCTION] OAuthClient: Creates a client to authenticate with the API
        * [FUNCTION] get_data: Gets data from the API
        * [FUNCTION] get_query_param: Gets query parameters from the URL
        * [FUNCTION] pageByID: Gets the ID from the URL and displays it on the page
        * [FUNCTION] jsonobj: Displays JSON data on the page, and refreshes 
        * [FUNCTION] NotificationWebHooks: Called when the API state changes
        * [FUNCTION] api_root: Displays the JSON data on the page
        * [FUNCTION] index: Displays the JSON data on the page
        * [FUNCTION] UPLOADFILEFORMS: Uploads a file to the server
        * [FUNCTION] api_web: Displays the JSON data on the page
    
    [IMPORTANT FLASK COMMANDS]
        * [render_template_string]: Renders a template from a string, rather than a file.
        * [render_template]: Renders a template from the template folder with the given context.
        * [blueprint]: A blueprint is a way to organize related views.
        * [make_response]: Creates a response object with the given status code and data.
        * [send_file]: Sends a file to the client.
        * [FileField, SubmitField ]: Creates a file upload field and a submit button.
        * [FlaskForm]: A class that represents a form.
        * [werkzeug.utils]: A collection of utilities for WSGI applications.
          -> wsgi applications are a standard interface between web servers and Python web applications or frameworks.
          that allow downloads and uploads of files.
          
          [flash]: Flashes a message to the client.
          
          
    [LIBRARIES USED]
        * Flask: A micro web framework written in Python.
            + SOCKETIO: A Flask extension that adds WebSocket support to your application.
        * Flask-SocketIO: A Flask extension that adds WebSocket support to your application.
        * Requests: A Python library for sending HTTP requests.
        * REDIRECT: Redirects the user to another endpoint.
        * URL_FOR: Generates a URL to the given endpoint with the method provided.
'''
'''
    REMEMBER: Web are reverse API's that react to events 
    To expose your local server to the internet, you can use ngrok.
    
    [ngrok] takes your localhost and gives you a public URL that you can use to expose your local web server to the internet.
    First install ngrok by following the instructions on their website: https://ngrok.com/download
    Then run ngrok (torun) ngrok http http://localhost:8080
    https://dashboard.ngrok.com/
    
    Then Start Flask: 
    export FLASK_APP=app
    flask run
    
    -- USE 'STAR' FEATURE ON SERVICECHANNEL API TO TEST HOOKS --- 
    
  
'''
class StoreData:
    def __init__(self):
        self.data = None

    def store(self, value):
        self.data = value

    def retrieve(self):
        return self.data

class StoreTokens:
    def __init__(self):
        self.token = None

    def store(self, value):
        self.token = value
    def retrieve(self):
        return self.token

class StorePrettyJson:
    def __init__(self):
        self.data = None

    def store(self, value):
        self.data = value

    def retrieve(self):
        return self.data



''' SUPPORT CLASS FOR UPLOAD FILE '''
class UPLOADFILEFORMS(FlaskForm):
    file = FileField('File', validators=[InputRequired()]) ## Validates that the upload is not empty
    submit = SubmitField('submit')



class JSONDocumentCreator:

    def __init__(self, json_data):
        self.data = json.loads(json_data)
        self.doc = Document()
        self.doc.add_heading('JSON Data Document', 0)

    def add_json_to_doc(self, data=None, level=0):

        print("[!] Adding JSON to document")
        if data is None:
            data = self.data

        if isinstance(data, dict):
            for key, value in data.items():
                self.doc.add_paragraph(f"{'  ' * level}{key}:", style='ListBullet')
                self.add_json_to_doc(value, level + 1)
        elif isinstance(data, list):
            for item in data:
                self.add_json_to_doc(item, level)
        else:
            self.doc.add_paragraph(f"{'  ' * level}{data}")

    def save_document(self, file_name):
        print("[!] Saving document")
        self.doc.save(file_name)

class TokenParser:
    def __init__(self, toParse):
        self.token_table = json.loads(toParse)
        self.access_token = self.token_table['access_token']
        self.refresh_token = self.token_table['refresh_token']

    def get_token_table(self):
        return self.token_table

    def get_access_token(self):
        return self.access_token

    def get_refresh_token(self):
        return self.refresh_token


''' BUILDS CUSTOM HEADERS FOR API REQUESTS '''
def makeCustomHeaders(token):
    myResponse = make_response('Response')
    myResponse.headers['customHeader'] = 'This is a custom header'
    myResponse.status_code = 403
    myResponse.mimetype = 'video/mp4'
    return myResponse


storedata = StoreData()
storetokens = StoreTokens()
storePrettyJson = StorePrettyJson()

class OAuthClient:
    def __init__(self, tokenURL, auth_header, username, password):
        self.tokenURL = tokenURL
        self.auth_header = auth_header
        self.username = username
        self.password = password
        self.access_token = None
        self.refresh_token = None

    def get_resp(self):
        headers = {
            'Authorization': self.auth_header,
            'Content-Type': "application/x-www-form-urlencoded"
        }

        payload = {
            "username": self.username,
            "password": self.password,
            "grant_type": "password"
        }
        print("sending reqeust to ", self.tokenURL)
        # Send the POST request
        response = requests.post(self.tokenURL, headers=headers, data=payload)
        print("[+]\n", response.text)
        print("[+]\n", response.json)
        print("[+]\n", response.status_code)
        if response.status_code == 200:
            new_tokens = response.json()
            new_token_str = json.dumps(new_tokens, indent=4)
            self.access_token = new_tokens['access_token']
            self.refresh_token = new_tokens['refresh_token']

           ### SENDS TOKENS TO TOKEN PARSER
            T = TokenParser(new_token_str)
            T.access_token = self.access_token
            T.refresh_token = self.refresh_token

           # storedata.store(response.json())
            storetokens.store(new_tokens)  ## Stores Data

            print("[+]\n", response.text)
            print('[+] TOKEN: \n', self.access_token)
            print('[+] REFRESH TOKEN: \n', self.refresh_token)
            print('[+] EXPIRES IN \n', new_tokens['expires_in'])
        else:
            print(f"Failed to authenticate: {response.status_code} - {response.content}")
            print(response.reason)

    def get_access_token(self):
        return self.access_token

    def get_refresh_token(self):
        return self.refresh_token


url = 'http://localhost:5000/webhook'
tokenURL = "https://sb2login.servicechannel.com/oauth/token"
auth_header = "Basic UFIuMjA5ODQ1NzU5OC4yNkFBMkVFOS0xMTUwLTRDQzctOEU0Qi1ENzk2MUM3NDFBRDY6RjQ4RDcxRDItNEFEOS00RjQ0LUE0MTktRjMxMUFBMUNEQTk3"
client = OAuthClient(tokenURL, auth_header, "fuad@dedicatedglass.com", "Dedicated!234")

client.get_resp()
access_token = client.get_access_token()
refresh_token = client.get_refresh_token()


print("[+] ACCESS TOKEN: ", access_token)
print("[+] REFRESH TOKEN: ", refresh_token)

print("[+] STARTING FLASK: ")
logging.basicConfig(level=logging.DEBUG)


''' THIS IS USED TO CREATE A BLUEPRINT TEMPLATE FOR ROOT URL '''
app = Flask(__name__)
socketio = SocketIO(app)

''' REGISTERING BLUEPRINTS '''
# app.register_blueprint(views, url_prefix="/web")


''' FUNCTION TO UPLOAD DOCUMENT 
    * [!] UPLOAD FOLDER: Folder to store the uploaded files
        [!] FORM DATA: form.file.data (FILE THATS UPLOADED) 
        [!] FILENAME: file.filename (NAME OF FILE)
        [!] FILE.SAVE: Saves the file to the server BY JOINING THE PATH to find ROOT DIR 
            - os.path.join: Joins one or more path components
            - os.path.abspath: Returns the absolute path of the file
            - secure_filename: Returns a secure version of the filename
            - InputRequired: Validates that the input is not empty (doesnt allow empty files)
            
    * [!] (Once Saved) RETURNS: File has been updated successfully
    
    [!] [UPLOAD_FOLDER]: app.config['UPLOAD_FOLDER'] (FOLDER TO UPLOAD FILE)
    [!] [SECRET_KEY]: app.config['SECRET_KEY'] (SECRET KEY FOR filesystem) 
'''


''' SECURITY MEASURES'
    -> Primarly for uploading files, to prevent malicious files from being uploaded
'''
app.config['SECRET_KEY'] = 'hello'
app.config['UPLOAD_FOLDER'] = 'static/files'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
UPLOAD_FOLDER = 'static/uploads'
app.secret_key = "secret key"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024


@app.route('/upload', methods=['GET', 'POST'])
def upload():
    print("[!] UPLOADING FILE\n\n")
    form = UPLOADFILEFORMS()
    if form.validate_on_submit():
        print("[+] FORM DATA: (VALIDATED) ", form.file.data)
        file = form.file.data
        file.save(os.path.join(os.path.abspath(os.path.dirname(__file__)),app.config['UPLOAD_FOLDER'],secure_filename(file.filename)))
        print("[!] FILE SAVED: ", file.filename)
        return f"file has been updated successfully \n [FORM DATA]{form.file.data} \n [!] FILENAME: {file.filename}"

    return render_template('index.html', CONTENT=storedata.retrieve(), form=form)


'''     
    FUNCTION TO SEND FILES FOR CLIENT DOWNLOAD 
        * sends webhook.json file as an attachment to client 
'''
@app.route('/download')
def download():
    return send_file("./webook.json", as_attachment=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def api_root():
    data = storedata.retrieve()
    # Convert the Python dictionary to a pretty-printed JSON string
    pretty_json = json.dumps(data, indent=4)
    print("[!] Data: \n", pretty_json)


    ''' LOGIC FOR UPLOADING PICTURES'''
    # if 'file' not in request.files:
    #     flash('select an image to upload')
    #     #return redirect(request.url)
    # file= request.files['file']
    # if file.filename == '':
    #     flash('No selected file')
    #     #return redirect(request.url)
    # if file.filename == '':
    #     flash('No selected file')
    #     #return redirect(request.url)
    # if file and allowed_file(file.filename):
    #     filename = secure_filename(file.filename)
    #     file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    #     flash('File successfully uploaded')
    #     #return redirect('/')
    #     return render_template('index.html', CONTENT=pretty_json, TOKEN=storetokens.retrieve(), form=form,
    #                            TIME=time.ctime())

    ''' This makes the JSON data more readable by adding indentation and line breaks 
        html.escape: Escapes special characters in the JSON string
    '''
    escaped_json = html.escape(pretty_json)
    html_string = f"""
                   <!DOCTYPE html>
               <html>
               <head>
                   <title>JSON Display</title>
               </head>
               <body>
                   <h1>JSON Data:</h1>
                   <pre>{escaped_json}</pre>
               </body>
               </html>
           """

    # Write the HTML string to a file
    with open("json_display.html", "w") as f:
        f.write(html_string)


    ''' LOGIC TO SETUP FORMS FOR UPLOADING FILES '''
    form = UPLOADFILEFORMS()
    if form.validate_on_submit():
        print("[+] FORM DATA: (VALIDATED) ", form.file.data)
        file = form.file.data
        file.save(os.path.join(os.path.abspath(os.path.dirname(__file__)), app.config['UPLOAD_FOLDER'],
                               secure_filename(file.filename)))

        print("[!] FILE SAVED: ", file.filename)
        return f"file has been updated successfully \n [FORM DATA]{form.file.data} \n [!] FILENAME: {file.filename}"


    return render_template('index.html', CONTENT=pretty_json, TOKEN=storetokens.retrieve(), form=form, TIME=time.ctime())

   # return f' \n[REFRESH TOKEN]{refresh_token}\n {html_string}'


''' HOOK IS CALLED WHEN API STATE CHANGES  '''
@app.route('/NotificationWebHooks', methods=['POST'])
def NotificationWebHooks():
    print("[!] NOTIFICATION WEBHOOKS\n\n")

    if request.method == 'POST':
        data = request.get_json()
        print("[!] Webhook Data\n", data)

        ## prettify and store data
        pp.pprint(data, compact=True)
        pretty_json_str = pp.pformat(data, compact=True).replace("'", '"')
        storedata.store(data)  ## Stores Data

       ## gets data from JSON, then writes to file, with unique ID
        id = data['Object']['Id']
        with open(f'inbound_wo{id}.json', 'w') as f:
            f.write(pretty_json_str)
        print("********************")

        # Convert the JSON string to a Python dictionary
     #   hashed_data = json.loads(data)
        print("++++++++++++++++++")

        return redirect(url_for('api_root'))
        #return 'success', 200
        # return jsonify({"status": "Success", "message": "Document generated successfully"}), 200
    client.get_resp()
    access_token_internal = client.get_access_token()
    refresh_token_internal = client.get_refresh_token()
    data = request.get_json()
    token = request.headers.get('Authorization')
    app.logger.error(access_token_internal)
    app.logger.error(refresh_token_internal)
    logging.debug("Access Token: %s", access_token_internal)
    logging.debug("Refresh Token: %s", refresh_token_internal)

    logging.debug(data)
    if not token:
        return redirect(url_for('api_root'))
      #  return jsonify({"status": "Error", "message": "No token provided"}), 401


    logging.debug("Webhook Data: %s", data)
    print("[!] Webhook Data\n", data)
    print("[!] Request Data \n", request.data)


    pp.pprint(data, compact=True) ## Pretty Print Data
    pretty_json_str = pp.pformat(data, compact=True).replace("'", '"')


    with open(f'webook.json', 'w') as f:
        f.write(pretty_json_str)

    if request.headers['Content-Type'] == 'application/json':
        data = request.json
        print("[!] Data: ", data)
        return f" DATA RECEIVED, DUMPING {json.dump(request.json)}"

    return jsonify({"status": "Success", "message": "Document generated successfully"}), 200


''' SIMPLE PAGE TO DISPLAYS WEBHOOK (HTTML)'''
@app.route('/web')
def index():
    print("[!] Accessing Index Page\n\n")
    token2 = OAuthClient.get_access_token
    token = storetokens.retrieve()
    data = storedata.retrieve()
    # Convert the Python dictionary to a pretty-printed JSON string
    pretty_json = json.dumps(data, indent=4)
    print("[!] Data: \n", pretty_json)
    # Escape the JSON string for HTML
    escaped_json = html.escape(pretty_json)

    storePrettyJson.store(escaped_json)

    print("[!] INDEX Data: \n", escaped_json)
    # creates a simple HTML page to display the JSON data
    html_string = f"""
            <!DOCTYPE html>
        <html>
        <head>
            <title>JSON Display</title>
        </head>
        <body>
            <h1>JSON Data:</h1>
            <h1>{escaped_json}</h1>
        </body>
        </html>
    """


    # Write the HTML string to a file
    with open("json_display.html", "w") as f:
        f.write(html_string)

    return render_template_string('''
    <!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta http-equiv="X-UA-Compatible" content="ie=edge">
    <title>Your Website Title</title>
    <link rel="stylesheet" href="style.css">
    <!-- Add your custom stylesheets above -->
    <script src="index.js"></script>
    <!-- Add your custom scripts above -->
</head>
<body>
    <header>
        <h1> {{ ACTION }}"</h1>
        <h2> {{ EVENTTYPE }}</h2>
        <h3> "TOKEN:" {{ TOKEN }} </h3>
    </header>
    <nav>
        <ul>

            <li> "Id:" {{ ID }}</li>
            <li> "LocationId:" {{ LOCATIONID }}</li>
            <li> "ProviderId:" {{ PROVIDERID }}</li>
            <li> "SubscriberId:" {{ SUBSCRIBERID }}</li>
            <li> "ACTION:" {{ ACTION }}</li>

        </ul>
    </nav>
    <main>
        <section>
            <h2> CONTENT </h2>
            <p>  {{ CONTENT }}</p>

<button id="downloadBtn">Download Document</button>

<script>
    document.getElementById('downloadBtn').addEventListener('click', function() {
        fetch('/download-document')
            .then(response => response.blob())
            .then(blob => {
                // Create a new URL for the blob
                const url = window.URL.createObjectURL(blob);
                // Create a link to download it
                const a = document.createElement('a');
                a.style.display = 'none';
                a.href = url;
                // Give the download a name
                a.download = 'document.pdf'; // Change the extension based on your document
                // Append the link to the body
                document.body.appendChild(a);
                // Trigger the download
                a.click();
                // Clean up
                window.URL.revokeObjectURL(url);
                document.body.removeChild(a);
            })
            .catch(error => console.error('Error:', error));
    });
</script>
           <pre>{{ escaped_json }}</pre>
        </section>
    </main>
    <footer>
        <p>&copy; 2019 Your Website Title</p>
    </footer>
</body>
</html>
    ''', TOKEN=token, CONTENT=escaped_json, ID=token2, LOCATIONID=token2)



    # return render_template('index.html',
    #                        TITLE="WEBHOOK",
    #                        CONTENT=pretty_json,
    #                        TOKEN= token
    #                        )


''' 
    CREATES A DYNAMIC WEBPAGE THAT DISPLAYS (AND CHANGES) BASED OFF OF WHAT IS INSIDE OF URL 
    EXAMPLE: http://localhost:5000/web/web/change/1234
    WILL DISPLAY: 1234 IN TITLE 
'''


@app.route('/web/change/<string:ID>')
def pageByID(ID):
    return render_template("index.html",
                           TITLE=ID,
                           CONTENT="This is a webhook")


'''
    TAKES JSON DATA FROM WEBHOOK AND DISPLAYS IT ON THE PAGE
    ** -->> CURRENTLY SET TO REFRESH EVERY 5 SECONDS
    * [<meta http-equiv="refresh" content="5">] (CHANGE TO 0 TO STOP REFRESH) *
    
    JSONIFY IS USED TO RETURN A PYTHON DICT TO JSON FORMAT
    EXAMPLE: http://localhost:5000/web/jsonobj/ 
    WILL DISPLAY: JSON DATA IN BODY 
'''

@app.route('/web/jsonobj')
def jsonobj():

    data2 = storedata.retrieve()
    d = storetokens.retrieve()
    # id = d['Object']['Id']
    # send_file(f'./inbound_wo{id}.json', as_attachment=True)

    #  return jsonify(data2)
    return render_template_string("""
        <html>
            <head>
                <meta http-equiv="refresh" content="0">
            </head>
            <body>
                <h1>Page refreshed at: {{ TIME }}</h1>
                <h2>JSON Data: {{ DATA2 }} </h2>
                
                
<button id="downloadBtn">Download Document</button>

<script>
    document.getElementById('downloadBtn').addEventListener('click', function() {
        fetch('/download-document')
            .then(response => response.blob())
            .then(blob => {
                // Create a new URL for the blob
                const url = window.URL.createObjectURL(blob);
                // Create a link to download it
                const a = document.createElement('a');
                a.style.display = 'none';
                a.href = url;
                // Give the download a name
                a.download = 'document.pdf'; // Change the extension based on your document
                // Append the link to the body
                document.body.appendChild(a);
                // Trigger the download
                a.click();
                // Clean up
                window.URL.revokeObjectURL(url);
                document.body.removeChild(a);
            })
            .catch(error => console.error('Error:', error));
    });
</script>
            </body>
        </html>
        """, TIME=time.ctime(), DATA2=data2, ESCAPED=storePrettyJson.retrieve())


def get_data():
    data = request.get_json()
    return jsonify(data)

''' 
    PASSES A QUERY [?] PARAMATER TO THE URL 
    IT PULLS THE QUERY PARAMATER AND DISPLAYS IT ON THE PAGE
    EXAMPLE: http://localhost:5000/web/query?name=FUAD --- *refer: [?]name=FUAD)
    WILL DISPLAY: FUAD IN TITLE

    USAGE: 'args' is a dictionary that contains all the query parameters in the URL
'''


@app.route('/search/query')
def get_query_param():
    args = request.args
    id = request.args.get('TOKEN')
    return render_template("index.html",
                           TITLE=id,
                           CONTENT="This is a webhook")


if __name__ == '__main__':
    app.run(host="localhost", port=5005, debug=True)
    socketio.run(app, port=5005, debug=True)